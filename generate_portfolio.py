import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import os

# Configuration
SOURCE_DIR = r"c:\Users\Trinity\Documents\trinity\ltdatkinson\LTD portfolio"
OUTPUT_DIR = r"c:\Users\Trinity\Documents\trinity\ltdatkinson\LTD portfolio"
SOURCE_FILES = [
    "Portfolio Badge Content Creation.docx",
    "Daydream Technology Badge Content Creation.docx"
]
AUTHOR = "Joshua Atkinson"
TITLE = "The Daydream Initiative: LDT Competency Portfolio"
INSTITUTION = "Purdue University"

def set_apa_style(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Paragraph formatting
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 2.0
    paragraph_format.space_after = Pt(0)
    
    # Margins (1 inch)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

def create_title_page(doc):
    # Page number top right
    header = doc.sections[0].header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    # Add page number field
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

    # Title content
    for _ in range(4): doc.add_paragraph() # Spacing
    
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run(TITLE)
    r_title.bold = True
    
    doc.add_paragraph() # Spacing
    
    p_info = doc.add_paragraph()
    p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_info.add_run(f"{AUTHOR}\n{INSTITUTION}\n")
    p_info.add_run(f"Learning Design and Technology\n")
    p_info.add_run(datetime.date.today().strftime("%B %d, %Y"))
    
    doc.add_page_break()

def process_source_file(doc, filename):
    path = os.path.join(SOURCE_DIR, filename)
    if not os.path.exists(path):
        print(f"Warning: File not found: {path}")
        return

    source_doc = docx.Document(path)
    
    # Add a section header for the file content
    doc.add_paragraph(filename.replace(".docx", ""), style='Heading 1')
    
    for paragraph in source_doc.paragraphs:
        if paragraph.text.strip():
            # Copy content
            new_p = doc.add_paragraph(paragraph.text)
            # Try to preserve basic styling (bold/italic) if possible, 
            # but for now just raw text to ensure APA font application
            # If it's a heading in source, map to heading in dest
            if paragraph.style.name.startswith('Heading'):
                new_p.style = paragraph.style.name

def main():
    doc = docx.Document()
    set_apa_style(doc)
    create_title_page(doc)
    
    for filename in SOURCE_FILES:
        process_source_file(doc, filename)
        doc.add_page_break()
        
    output_filename = f"Master_Bible_Portfolio_{datetime.date.today()}.docx"
    output_path = os.path.join(OUTPUT_DIR, output_filename)
    doc.save(output_path)
    print(f"Created: {output_path}")

    # Generate JSON for Website
    import json
    portfolio_data = {
        "owner": AUTHOR,
        "title": TITLE,
        "categories": [
            {"id": "foundations", "title": "Professional Foundations", "description": "Core competencies in instructional design research and theory."},
            {"id": "planning", "title": "Planning and Analysis", "description": "Skills in analyzing learning needs and planning instructional interventions."},
            {"id": "design", "title": "Design and Development", "description": "Creating engaging and effective learning experiences."},
            {"id": "evaluation", "title": "Evaluation and Implementation", "description": "Assessing learning outcomes and implementing solutions."}
        ],
        "badges": []
    }
    
    # Simple extraction logic (this would need to be more robust in a real parser, 
    # but for now we will create a placeholder structure populated with the raw text chunks 
    # so the user can see the content in the website and we can refine the parsing later)
    
    # For the purpose of this task, we will create one "Mega Badge" per source file 
    # to demonstrate the data flow, as parsing the specific "Challenge 1", "Challenge 2" 
    # structure from the raw docx without strict delimiters is error-prone.
    # We will ask the user to refine the JSON or we can refine the parser if strict delimiters exist.
    
    for filename in SOURCE_FILES:
        path = os.path.join(SOURCE_DIR, filename)
        if os.path.exists(path):
            source_doc = docx.Document(path)
            full_text = "\n".join([p.text for p in source_doc.paragraphs if p.text.strip()])
            
            badge = {
                "id": filename.replace(" ", "_").replace(".docx", "").lower(),
                "categoryId": "design", # Defaulting to design for now
                "title": filename.replace(".docx", ""),
                "image": "/assets/badges/default.png",
                "artifacts": [
                    {
                        "title": "Portfolio Artifact",
                        "origin": "Derived from " + filename,
                        "summary": full_text[:500] + "...", # Preview
                        "reflection": {
                            "challenge": "Complete the requirements for " + filename,
                            "content": full_text[:1000], # First 1000 chars as placeholder content
                            "competency_alignment": "Demonstrated skills in " + filename
                        },
                        "full_content": full_text # Store full text for detail view
                    }
                ]
            }
            portfolio_data["badges"].append(badge)

    json_output_path = r"c:\Users\Trinity\Documents\trinity\ltdatkinson\client\src\data\portfolioData.json"
    with open(json_output_path, "w") as f:
        json.dump(portfolio_data, f, indent=2)
    print(f"Created: {json_output_path}")

if __name__ == "__main__":
    main()
